#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""完整转换docx为markdown - 区分目录和正文"""

import docx
import re
import os

def convert_docx_complete(docx_path):
    """完整转换docx为markdown"""

    doc = docx.Document(docx_path)
    markdown_content = []

    print(f"文档信息: {len(doc.paragraphs)} 段落, {len(doc.tables)} 表格")

    # 按文档顺序处理
    body_elements = list(doc.element.body)

    toc_started = False
    toc_ended = False
    content_started = False

    for element in body_elements:
        if element.tag.endswith('p'):
            # 找到对应的段落对象
            para = None
            for p in doc.paragraphs:
                if p._element == element:
                    para = p
                    break

            if para:
                style_name = para.style.name if para.style else ''
                text = para.text.strip()

                # 检测目录开始（"目    录"）
                if '目' in text and '录' in text and len(text) < 10:
                    toc_started = True
                    markdown_content.append('## 目录')
                    continue

                # 检测目录结束（第一个正文标题"概述"）
                if '概述' == text or '第一章' in text or '第1章' in text:
                    toc_started = False
                    toc_ended = True
                    content_started = True

                # 目录条目处理
                if toc_started and not toc_ended:
                    if text:
                        # 目录条目保留原格式，但标记为目录
                        # 格式：章节标题 + TAB + 页码
                        # 转换为：章节标题（页码）
                        parts = text.split('\t')
                        if len(parts) >= 2:
                            # 去除页码，保留章节名称
                            chapter_name = parts[0]
                            page_num = parts[-1] if parts[-1].isdigit() else ''
                            if page_num:
                                markdown_content.append(f"- {chapter_name}（第{page_num}页）")
                            else:
                                markdown_content.append(f"- {text}")
                        else:
                            markdown_content.append(f"- {text}")
                    continue

                # 正文内容处理
                if content_started or 'Heading' in style_name or '标题' in style_name:
                    processed = process_paragraph(para)
                    if processed:
                        markdown_content.append(processed)

        elif element.tag.endswith('tbl'):
            # 处理表格
            table = None
            for t in doc.tables:
                if t._element == element:
                    table = t
                    break

            if table:
                table_md = process_table(table)
                if table_md:
                    markdown_content.append(table_md)

    return '\n\n'.join(markdown_content)

def process_paragraph(para):
    """处理单个段落"""
    text = para.text.strip()
    if not text:
        return ''

    style_name = para.style.name if para.style else ''

    # 判断标题级别
    level = 0

    # 通过样式名称判断
    if 'Heading 1' in style_name or '标题 1' in style_name:
        level = 1
    elif 'Heading 2' in style_name or '标题 2' in style_name:
        level = 2
    elif 'Heading 3' in style_name or '标题 3' in style_name:
        level = 3
    elif 'Heading 4' in style_name or '标题 4' in style_name:
        level = 4
    elif 'Heading 5' in style_name or '标题 5' in style_name:
        level = 5
    elif 'Heading 6' in style_name or '标题 6' in style_name:
        level = 6

    # 处理格式
    formatted_text = process_formatting(para)

    if level > 0:
        return '#' * level + ' ' + formatted_text
    else:
        # 检查是否是列表项
        if is_list_item(para):
            return process_list_item(para, formatted_text)
        return formatted_text

def process_formatting(para):
    """处理段落中的格式"""
    result = []

    for run in para.runs:
        text = run.text

        if not text:
            continue

        # 处理粗体
        if run.bold:
            text = '**' + text + '**'

        # 处理斜体
        if run.italic:
            text = '*' + text + '*'

        # 处理下划线
        if run.underline:
            text = '<u>' + text + '</u>'

        result.append(text)

    return ''.join(result)

def is_list_item(para):
    """判断是否是列表项"""
    style_name = para.style.name if para.style else ''

    # 通过样式判断
    list_styles = ['List', '列表', 'ListNumber', 'ListBullet', '编号', '项目符号']
    for ls in list_styles:
        if ls in style_name:
            return True

    # 通过文本内容判断
    text = para.text.strip()
    if re.match(r'^\d+\.\s', text):  # 有序列表：1. xxx
        return True
    if re.match(r'^[-•]\s', text):  # 无序列表：- xxx 或 • xxx
        return True
    if re.match(r'^\（\d+\）', text):  # 中文有序列表：（1）xxx
        return True

    return False

def process_list_item(para, formatted_text):
    """处理列表项"""
    text = para.text.strip()

    # 有序列表 - 数字开头
    if re.match(r'^\d+\.\s', text):
        return formatted_text  # 保持原格式

    # 中文有序列表
    if re.match(r'^\（\d+\）', text):
        return formatted_text  # 保持原格式

    # 无序列表
    if re.match(r'^[-•]\s', text):
        return '- ' + text[2:] if len(text) > 2 else formatted_text

    return formatted_text

def process_table(table):
    """处理表格"""
    if not table.rows:
        return ''

    markdown_rows = []

    # 处理所有行
    for row_idx, row in enumerate(table.rows):
        cells = []
        for cell in row.cells:
            # 获取单元格文本
            cell_text = get_cell_text(cell)
            cells.append(cell_text)

        if cells:
            markdown_rows.append('| ' + ' | '.join(cells) + ' |')

            # 如果是第一行，添加分隔线
            if row_idx == 0:
                separator = '| ' + ' | '.join(['---' for _ in cells]) + ' |'
                markdown_rows.append(separator)

    return '\n'.join(markdown_rows)

def get_cell_text(cell):
    """获取单元格文本"""
    texts = []
    for para in cell.paragraphs:
        text = para.text.strip()
        if text:
            # 处理单元格内的格式
            formatted = process_formatting(para)
            texts.append(formatted)

    cell_text = '\n'.join(texts)
    # 替换换行为空格（表格单元格内不支持换行）
    cell_text = cell_text.replace('\n', ' ')
    return cell_text

def main():
    docx_file = 'IT项目成本估算监控系统.docx'
    output_file = 'IT项目成本估算监控系统.md'

    try:
        print("开始完整转换（区分目录和正文）...")
        markdown_content = convert_docx_complete(docx_file)

        # 添加文档元信息头部
        header = """---
title: IT项目成本估算监控系统
subtitle: 用户需求说明书
author: 胡月
date: 2026年3月30日
---

**IT项目成本估算监控系统**

**用户需求说明书**

**制作单位：(部门)**

**文档版本号:**

**日期：**

**编写人员:胡月**

**校对人员:**

**2026 年 3 月 30 日**

**文档更改记录表**

**A** – 添加的  **M** – 修改的  **D** – 删除的

"""

        full_content = header + markdown_content

        # 保存到文件
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(full_content)

        print(f'\n转换成功！文件已保存到: {output_file}')
        print(f'文件大小: {len(full_content)} 字符')
        print(f'内容行数: {len(full_content.splitlines())} 行')

        # 验证关键内容
        keywords = ['概述', '业务概述', '功能描述', '数据要求', '非功能', '参考资料']
        print('\n=== 内容验证 ===')
        for kw in keywords:
            if kw in full_content:
                print(f'OK 包含 "{kw}"')
            else:
                print(f'MISS 缺失 "{kw}"')

        # 统计表格数量
        table_count = full_content.count('| --- |')
        print(f'\n转换表格数量: {table_count}')

    except Exception as e:
        print(f'转换失败: {str(e)}')
        import traceback
        traceback.print_exc()

if __name__ == '__main__':
    main()