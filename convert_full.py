#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""使用python-docx完整转换docx为markdown - 确保所有内容都保留"""

import docx
import re
import os

def convert_docx_to_markdown_full(docx_path):
    """完整转换docx为markdown"""

    doc = docx.Document(docx_path)
    markdown_content = []

    print(f"文档信息: {len(doc.paragraphs)} 段落, {len(doc.tables)} 表格")

    # 首先处理所有段落
    para_idx = 0
    table_idx = 0

    # 需要跟踪表格在文档中的位置
    # python-docx中表格和段落是分开存储的，需要按文档顺序处理
    # 使用element来获取文档顺序

    body_elements = list(doc.element.body)

    for element in body_elements:
        if element.tag.endswith('p'):
            # 处理段落
            para = None
            for p in doc.paragraphs:
                if p._element == element:
                    para = p
                    break

            if para:
                text = process_paragraph(para)
                if text:
                    markdown_content.append(text)

        elif element.tag.endswith('tbl'):
            # 处理表格
            table = None
            for t in doc.tables:
                if t._element == element:
                    table = t
                    break

            if table:
                table_md = process_table(table)
                if table_md:
                    markdown_content.append(table_md)

    return '\n\n'.join(markdown_content)

def process_paragraph(para):
    """处理单个段落"""
    text = para.text.strip()
    if not text:
        return ''

    style_name = para.style.name if para.style else ''

    # 判断标题级别
    level = 0

    # 通过样式名称判断
    if 'Heading 1' in style_name or '标题 1' in style_name:
        level = 1
    elif 'Heading 2' in style_name or '标题 2' in style_name:
        level = 2
    elif 'Heading 3' in style_name or '标题 3' in style_name:
        level = 3
    elif 'Heading 4' in style_name or '标题 4' in style_name:
        level = 4
    elif 'Heading 5' in style_name or '标题 5' in style_name:
        level = 5
    elif 'Heading 6' in style_name or '标题 6' in style_name:
        level = 6

    # 通过文本内容判断章节标题
    if level == 0:
        if text.startswith('第') and '章' in text:
            level = 1
        elif re.match(r'^\d+\.\d+\s', text):  # 如 "1.1 需求背景"
            level = 2
        elif re.match(r'^\d+\.\d+\.\d+\s', text):  # 如 "2.3.1 我行及同业现状"
            level = 3
        elif re.match(r'^\d+\s', text) and len(text) < 50:  # 简短数字开头
            # 可能是列表项，不作为标题
            pass

    # 处理粗体、斜体等格式
    formatted_text = process_formatting(para)

    if level > 0:
        return '#' * level + ' ' + formatted_text
    else:
        # 检查是否是列表项
        if is_list_item(para):
            return process_list_item(para, formatted_text)
        return formatted_text

def process_formatting(para):
    """处理段落中的格式"""
    result = []

    for run in para.runs:
        text = run.text

        if not text:
            continue

        # 处理粗体
        if run.bold:
            text = '**' + text + '**'

        # 处理斜体
        if run.italic:
            text = '*' + text + '*'

        # 处理下划线
        if run.underline:
            text = '<u>' + text + '</u>'

        result.append(text)

    return ''.join(result)

def is_list_item(para):
    """判断是否是列表项"""
    style_name = para.style.name if para.style else ''

    # 通过样式判断
    list_styles = ['List', '列表', 'ListNumber', 'ListBullet', '编号', '项目符号']
    for ls in list_styles:
        if ls in style_name:
            return True

    # 通过文本内容判断
    text = para.text.strip()
    if re.match(r'^\d+\.\s', text):  # 有序列表：1. xxx
        return True
    if re.match(r'^[-•]\s', text):  # 无序列表：- xxx 或 • xxx
        return True
    if re.match(r'^\（\d+\）', text):  # 中文有序列表：（1）xxx
        return True
    if re.match(r'^[（(][一二三四五六七八九十]+[）)]', text):  # 中文列表
        return True

    return False

def process_list_item(para, formatted_text):
    """处理列表项"""
    text = para.text.strip()

    # 有序列表 - 数字开头
    if re.match(r'^\d+\.\s', text):
        return formatted_text  # 保持原格式

    # 中文有序列表
    if re.match(r'^\（\d+\）', text):
        return formatted_text  # 保持原格式

    # 无序列表
    if re.match(r'^[-•]\s', text):
        return '- ' + text[2:] if len(text) > 2 else formatted_text

    return formatted_text

def process_table(table):
    """处理表格"""
    if not table.rows:
        return ''

    markdown_rows = []

    # 处理所有行
    for row_idx, row in enumerate(table.rows):
        cells = []
        for cell in row.cells:
            # 获取单元格文本
            cell_text = get_cell_text(cell)
            cells.append(cell_text)

        if cells:
            markdown_rows.append('| ' + ' | '.join(cells) + ' |')

            # 如果是第一行，添加分隔线
            if row_idx == 0:
                separator = '| ' + ' | '.join(['---' for _ in cells]) + ' |'
                markdown_rows.append(separator)

    return '\n'.join(markdown_rows)

def get_cell_text(cell):
    """获取单元格文本"""
    texts = []
    for para in cell.paragraphs:
        text = para.text.strip()
        if text:
            texts.append(text)

    cell_text = '\n'.join(texts)
    # 替换换行为空格（表格单元格内不支持换行）
    cell_text = cell_text.replace('\n', ' ')
    return cell_text

def main():
    docx_file = 'IT项目成本估算监控系统.docx'
    output_file = 'IT项目成本估算监控系统.md'

    try:
        print("开始完整转换...")
        markdown_content = convert_docx_to_markdown_full(docx_file)

        # 添加文档元信息头部
        header = """---
title: IT项目成本估算监控系统
subtitle: 用户需求说明书
author: 胡月
date: 2026年3月30日
---

"""

        full_content = header + markdown_content

        # 保存到文件
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(full_content)

        print(f'\n转换成功！文件已保存到: {output_file}')
        print(f'文件大小: {len(full_content)} 字符')
        print(f'内容行数: {len(full_content.splitlines())} 行')

        # 验证关键内容
        keywords = ['第4章', '第5章', '参考资料', '数据要求', '非功能']
        print('\n=== 内容验证 ===')
        for kw in keywords:
            if kw in full_content:
                print(f'✓ 包含 "{kw}"')
            else:
                print(f'✗ 缺失 "{kw}"')

    except Exception as e:
        print(f'转换失败: {str(e)}')
        import traceback
        traceback.print_exc()

if __name__ == '__main__':
    main()