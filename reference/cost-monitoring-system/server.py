#!/usr/bin/env python3
"""
项目成本估算监控系统 v1.0 - 后端服务
功能模块：
1. 登录认证
2. 首页仪表盘
3. 工作量评估（集成已有模块）
4. 成本消耗预估
5. 成本偏差分析

运行方式: python3 server.py
访问地址: http://localhost:8502
"""

import os
import json
import io
import re
import time
import tempfile
import subprocess
import traceback
from pathlib import Path
from collections import OrderedDict

import tornado.ioloop
import tornado.web
import requests
import openpyxl
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# ===================== 配置 =====================

PORT = 8501
OUTPUT_DIR = Path(__file__).parent

# 大模型接口
LLM_API_URL = 'https://www.finna.com.cn/v1/chat/completions'
LLM_MODEL   = 'qwq-32b'
API_KEY     = 'app-PvoiFWuSXcN4kwCBuplgOnnC'

# ── 新公式阶段定义（v2.1：严格对齐需求方案公式，仅含6个计算阶段） ──
# uses_tech_stack: 是否乘以技术栈难度系数
# 注意：需求方案公式只定义了以下6个阶段，业务测试和准生产不在公式中
PHASES_V2 = [
    {'key': 'requirements', 'name': '需求',    'uses_tech_stack': False},
    {'key': 'ui_design',    'name': 'UI设计',  'uses_tech_stack': False},
    {'key': 'tech_design',  'name': '技术设计','uses_tech_stack': True },
    {'key': 'development',  'name': '开发',    'uses_tech_stack': True },
    {'key': 'tech_testing', 'name': '技术测试','uses_tech_stack': True },
    {'key': 'perf_testing', 'name': '性能测试','uses_tech_stack': False},
    # 投产上线 = Σ(需求+设计+研发+技测+性测) × 2%，总计算后追加
]

# 阶段合规区间（占总人天百分比）
COMPLIANCE_RANGES = {
    '需求':    (12, 18),
    '设计':    (12,  20),   # UI设计+技术设计合并
    '开发':    (30, 40),
    '技术测试': (15, 25),
    '性能测试': (0, 12),
    # 投产上线固定 2%
    '投产上线': (2,  2),
}

# 各阶段颜色
PHASE_BG = {
    '需求':    'EBF3FB', 'UI设计':  'FFF9C4', '技术设计': 'E3F2FD',
    '开发':    'E2EFDA', '技术测试': 'FCE4D6', '性能测试': 'EDEDED',
    '业务测试': 'FFF2CC', '准生产':  'E8F5E9', '投产上线': 'F3E5F5',
}

# ── 默认参数 ──
DEFAULT_PARAMS = {
    'complexity_base': {
        'very_basic': 0.5,
        'basic': 1.0,
        'medium': 1.5,
        'complex': 2.0,
        'very_complex': 2.5
    },
    'association_coeff': 1.0,
    'flow_coeffs': {
        'requirements': 0.7,
        'ui_design':    0.3,
        'tech_design':  0.5,
        'development':  1.2,
        'tech_testing': 0.7,
        'perf_testing': 0.3,
    },
    'tech_stack_coeff': 1.3,
    'mgmt_coeff': 0.15,
    'daily_rates': {
        'product_manager': 2000,
        'ui_designer':     1800,
        'frontend_dev':    1800,
        'backend_dev':     2000,
        'func_tester':     1500,
        'perf_tester':     2000,
        'project_manager': 2000,
    },
    'go_live_pct': 0.02,   # 投产上线 = 前置阶段合计 × 2%
}

# ── 模拟用户数据 ──
USERS = {
    'admin': {'password': 'admin123', 'role': 'admin'},
    'user': {'password': 'user123', 'role': 'user'},
}

# ── 模拟项目数据 ──
PROJECTS = [
    {
        'id': 'PRJ-2026-001',
        'name': '人工智能Agent能力价值释放项目',
        'system_name': '成本预估智能平台',
        'status': '进行中',
        'estimated_cost': 1200000,
        'actual_cost': 950000,
        'start_date': '2026-01-15',
        'end_date': '2026-06-30',
        'manager': '张三',
        'progress': 65
    },
    {
        'id': 'PRJ-2026-002',
        'name': '运营数字化经营管理系统',
        'system_name': '运营管理平台',
        'status': '已完成',
        'estimated_cost': 800000,
        'actual_cost': 780000,
        'start_date': '2025-10-01',
        'end_date': '2026-01-31',
        'manager': '李四',
        'progress': 100
    },
    {
        'id': 'PRJ-2026-003',
        'name': '数字员工智能助手',
        'system_name': '数字员工平台',
        'status': '规划中',
        'estimated_cost': 1500000,
        'actual_cost': 0,
        'start_date': '2026-03-01',
        'end_date': '2026-09-30',
        'manager': '王五',
        'progress': 10
    }
]

# ===================== 文档提取 =====================

def extract_docx(path):
    if not HAS_DOCX:
        return ""
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return '\n'.join(parts)


def extract_doc(path):
    abs_path = os.path.abspath(path)
    # 方法1: 尝试用 LibreOffice 转换
    lo_candidates = ['libreoffice', 'soffice', '/usr/bin/libreoffice', '/usr/bin/soffice',
                     '/snap/bin/libreoffice']
    for lo in lo_candidates:
        try:
            with tempfile.TemporaryDirectory() as tmp:
                r = subprocess.run(
                    [lo, '--headless', '--convert-to', 'docx', '--outdir', tmp, abs_path],
                    capture_output=True, timeout=60
                )
                print(f"[doc->docx] {lo} returncode={r.returncode} stdout={r.stdout[:200]} stderr={r.stderr[:200]}")
                if r.returncode == 0:
                    # 查找转换后的文件（文件名可能有变化）
                    converted_files = list(Path(tmp).glob('*.docx'))
                    if converted_files:
                        text = extract_docx(str(converted_files[0]))
                        if text.strip():
                            print(f"[doc->docx] 转换成功, 提取文本长度={len(text)}")
                            return text
                    else:
                        print(f"[doc->docx] 转换后未找到.docx文件, 目录内容: {list(Path(tmp).iterdir())}")
        except FileNotFoundError:
            print(f"[doc->docx] {lo} 未安装")
            continue
        except Exception as e:
            print(f"[doc->docx] {lo} 异常: {e}")
            continue
    # 方法2: 尝试用 antiword
    try:
        r = subprocess.run(['antiword', abs_path], capture_output=True, timeout=30)
        if r.returncode == 0:
            text = r.stdout.decode('utf-8', errors='ignore').strip()
            if len(text) > 50:
                print(f"[doc] antiword 成功, 文本长度={len(text)}")
                return text
    except FileNotFoundError:
        print("[doc] antiword 未安装")
    except Exception as e:
        print(f"[doc] antiword 异常: {e}")
    # 方法3: 原始二进制尝试（最后手段，效果有限）
    try:
        raw = open(abs_path, 'rb').read()
        text = raw.decode('utf-16-le', errors='ignore')
        cleaned = re.sub(r'[^\u4e00-\u9fff\u0020-\u007e\n\r]+', ' ', text)
        cleaned = re.sub(r' {3,}', '\n', cleaned).strip()
        if len(cleaned) > 100:
            return cleaned
    except Exception:
        pass
    return ""


def extract_text(file_bytes, filename):
    suffix = Path(filename).suffix.lower()
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as f:
        f.write(file_bytes)
        tmp = f.name
    try:
        if suffix == '.docx':
            return extract_docx(tmp)
        elif suffix == '.doc':
            text = extract_doc(tmp)
            if not text.strip():
                raise ValueError(
                    '.doc格式读取失败（服务器缺少LibreOffice或antiword组件）。\n'
                    '解决方案：请将文件用Word另存为 .docx 格式后重新上传，或直接粘贴需求文本。'
                )
            return text
        elif suffix == '.txt':
            return file_bytes.decode('utf-8', errors='ignore')
        return ""
    finally:
        try:
            os.unlink(tmp)
        except Exception:
            pass

# ===================== AI 调用 =====================

def _llm_error_detail(res_json, raw_text):
    """从非标准/错误响应体中提取可读说明（部分网关 HTTP 200 但 body 无 choices）。"""
    if isinstance(res_json, dict):
        if res_json.get('message'):
            return str(res_json['message'])[:800]
        err = res_json.get('error')
        if isinstance(err, dict) and err.get('message'):
            return str(err['message'])[:800]
        if isinstance(err, str):
            return err[:800]
        if res_json.get('detail'):
            return str(res_json['detail'])[:800]
    return (raw_text or '')[:500]


def _consume_sse_text(resp):
    parts = []
    for raw_line in resp.iter_lines(decode_unicode=True):
        if not raw_line:
            continue
        line = raw_line.strip()
        if not line.startswith('data:'):
            continue
        data = line[5:].strip()
        if data == '[DONE]':
            break
        try:
            chunk = json.loads(data)
        except json.JSONDecodeError:
            continue
        err = chunk.get('error')
        if err:
            if isinstance(err, dict):
                raise ValueError(f"AI流式调用失败: {err.get('message', str(err))}")
            raise ValueError(f"AI流式调用失败: {err}")
        for choice in chunk.get('choices') or []:
            delta = choice.get('delta') or {}
            content = delta.get('content')
            if content:
                parts.append(content)
    return ''.join(parts)


def call_llm(prompt):
    payload = {
        "model": LLM_MODEL,
        "temperature": 0.7,
        "stream": True,
        "messages": [
            {"role": "system", "content": "你是一个专业的软件项目工作量评估专家，请只返回JSON格式结果，不要返回任何markdown标记或额外说明。"},
            {"role": "user", "content": prompt}
        ]
    }
    headers = {
        "Authorization": f"Bearer {API_KEY}",
        "Content-Type": "application/json; charset=utf-8"
    }
    with requests.post(LLM_API_URL, headers=headers, json=payload, timeout=300, stream=True) as resp:
        if resp.status_code != 200:
            try:
                res_json = resp.json()
            except Exception:
                res_json = {}
            detail = _llm_error_detail(res_json, resp.text)
            hint = ''
            low = detail.lower()
            if 'balance is insufficient' in low or '30001' in detail:
                hint = '（大模型账户余额不足或已欠费，请充值或更换 API Key 后重试）'
            raise ValueError(f'AI接口返回错误 (HTTP {resp.status_code}): {detail}{hint}')

        text = _consume_sse_text(resp)

    if text.strip():
        return text
    raise ValueError('AI流式返回为空或无法解析（请检查模型与网关 SSE 格式）')


def analyze_requirements(doc_text, api_key, management_coef="0.15"):
    """模型只判断需求难易程度（五档复杂度），不计算工作量。工作量由代码根据系数计算。"""
    prompt = f"""分析以下用户需求说明书，完成两项任务：

一、提取结构
1. 明确区分"模块(module)"与"功能点(function)"：
   - 模块：对若干相关业务功能的归类和分组，如"统一 AI 工作台""用户与权限管理"等。
   - 功能点：可独立开发和交付的最小业务能力单元，如"统一入口首页布局设计""登录鉴权与单点登录集成"等。
2. 文档中以「X.X.X.X 业务功能：XXX」或类似格式标注的，每一条视为一个独立功能点。
3. 先梳理出 3~15 个功能模块(module)，再将每个功能点(function)归入最合适的模块。若文档无显式模块名，可根据业务含义命名。
4. 从文档中提取项目名称和系统名称（如有）。

二、仅判断每个功能点的难易程度（五档复杂度）
请结合功能实现难度、业务规则复杂性、跨系统协同复杂度、异常分支/边界场景、数据处理与技术实现复杂度进行综合判断。
不要使用"按操作步骤数"作为判定规则。
每个功能点只输出以下五种之一：
   - very_basic（较为基础）
   - basic（基础）
   - medium（中等）
   - complex（复杂）
   - very_complex（极复杂）

三、为每个功能点给出"系统关联数量"建议（用于计算关联度系数）
结合需求描述判断该功能需要对接/依赖的外部系统数量（如：单点登录、主数据、财务、工单、消息、权限、流程等），只输出以下四档之一：
   - 1：仅本系统或只对接 1 个外部系统
   - 3：对接 2-5 个外部系统
   - 6：对接 5+ 个外部系统（用 6 表示）
若无法判断，默认输出 1。

不要计算任何系数或工作量，只输出模块、功能列表及每个功能的 complexity 和 association_systems。

请严格以 JSON 格式返回（不要包含任何其他说明文字，只返回 JSON）：
{{
  "project_name": "项目名称（未找到则空字符串）",
  "system_name": "系统名称（未找到则空字符串）",
  "modules": [
    {{
      "name": "模块名称",
      "functions": [
        {{"name": "功能名称", "complexity": "very_basic|basic|medium|complex|very_complex", "association_systems": 1}}
      ]
    }}
  ]
}}


需求文档内容（限前12000字）：
{doc_text[:12000]}"""

    text = call_llm(prompt)
    if not text:
        raise ValueError('AI模型返回为空')

    m = re.search(r'\{[\s\S]*\}', text)
    json_str = m.group() if m else text

    try:
        result = json.loads(json_str)
    except json.JSONDecodeError as e:
        raise ValueError(f'AI返回无法解析为JSON: {str(e)}。原始内容: {text[:200]}')

    result.setdefault('project_name', '')
    result.setdefault('system_name', '')
    result.setdefault('tech_stack', '微服务架构')
    result.setdefault('association_systems', 1)
    result.setdefault('modules', [])
    # 兼容：补齐每个功能点的 association_systems
    for m in result.get('modules', []) or []:
        for f in m.get('functions', []) or []:
            if 'association_systems' not in f:
                f['association_systems'] = result.get('association_systems', 1)
            try:
                f['association_systems'] = int(f.get('association_systems') or 1)
            except Exception:
                f['association_systems'] = 1
    return result

# ===================== 新版工作量计算引擎 =====================

def get_association_coeff(n_systems):
    """根据关联系统数量返回系数"""
    n = int(n_systems)
    if n <= 1:   return 1.0
    if n < 3:    return 1.5
    if n <= 5:   return 2.0
    return 3.0


def get_tech_stack_coeff(tech_stack_name):
    mapping = {
        '常规技术栈': 1.0,
        '微服务架构': 1.3,
        'AI中台':    1.4,
        '分布式事务': 1.6,
    }
    return mapping.get(tech_stack_name, 1.3)


def round_workload(v):
    """四舍五入到小数点后两位人天（保留两位小数），无最小值限制"""
    if v <= 0:
        return 0.0
    return round(v, 2)


def round_phase_total(v):
    """阶段合计四舍五入到0.5人天"""
    if v <= 0:
        return 0.0
    return round(round(v * 2) / 2, 1)


def estimate_workload_v2(modules, params=None):
    """
    新版工作量计算（v2.1 — 先求和再取整，消除逐功能点最小值膨胀）：
    公式：阶段工作量 = Σ(基准 × 关联度 × 流程系数 × [技术栈]) × (1 + 管理系数)
    返回: items列表 + 各阶段汇总 + 团队成本 + 合规结果 + 计算轨迹
    """
    if params is None:
        params = DEFAULT_PARAMS

    cb     = params.get('complexity_base', DEFAULT_PARAMS['complexity_base'])
    assoc_default  = float(params.get('association_coeff', DEFAULT_PARAMS['association_coeff']))
    fc     = params.get('flow_coeffs', DEFAULT_PARAMS['flow_coeffs'])
    tsc    = float(params.get('tech_stack_coeff', DEFAULT_PARAMS['tech_stack_coeff']))
    mgmt   = float(params.get('mgmt_coeff', DEFAULT_PARAMS['mgmt_coeff']))
    go_pct = float(params.get('go_live_pct', 0.02))

    items = []
    traces = []

    # 阶段原始值累加器（未乘管理系数、未取整）
    phase_raw_sums = OrderedDict()
    for ph in PHASES_V2:
        phase_raw_sums[ph['name']] = 0.0

    for module in modules:
        for func in module.get('functions', []):
            cx = str(func.get('complexity', 'medium')).lower()
            if cx in ('simple',):
                cx = 'basic'
            if cx not in ('very_basic', 'basic', 'medium', 'complex', 'very_complex'):
                cx = 'medium'

            base = cb.get(cx, cb.get('medium', 1.5))
            # 关联度系数：应优先取“审核功能列表”里每个功能点的关联系统数量换算出的系数
            assoc = assoc_default
            assoc_source = 'params.association_coeff'
            assoc_systems = None
            if isinstance(func, dict):
                if 'association_coeff' in func and func.get('association_coeff') is not None:
                    try:
                        assoc = float(func.get('association_coeff'))
                        assoc_source = 'func.association_coeff'
                    except Exception:
                        assoc = assoc_default
                elif 'association_systems' in func and func.get('association_systems') is not None:
                    try:
                        assoc_systems = int(func.get('association_systems'))
                        assoc = get_association_coeff(assoc_systems)
                        assoc_source = 'func.association_systems'
                    except Exception:
                        assoc = assoc_default
            trace = {
                'module': module['name'],
                'function': func['name'],
                'complexity': cx,
                'base': base,
                'assoc': assoc,
                'assoc_systems': assoc_systems,
                'assoc_source': assoc_source,
                'tech_stack': tsc,
                'mgmt': mgmt,
                'phases': {}
            }

            for ph in PHASES_V2:
                phase_key  = ph['key']
                phase_name = ph['name']
                flow = float(fc.get(phase_key, 0))

                if flow == 0:
                    raw_no_mgmt = 0.0
                elif ph['uses_tech_stack']:
                    raw_no_mgmt = base * assoc * flow * tsc
                else:
                    raw_no_mgmt = base * assoc * flow

                raw_with_mgmt = raw_no_mgmt * (1 + mgmt)
                # 单功能点工作量：保留两位小数（用于Excel逐行显示）
                wl = round_workload(raw_with_mgmt)
                # 日志跟踪：便于定位每个功能点/阶段的工作量来源
                if wl > 0:
                    try:
                        print(
                            f"[wl] module={module['name']} function={func['name']} "
                            f"phase={phase_name} raw_with_mgmt={raw_with_mgmt:.4f} wl={wl}"
                        )
                    except Exception:
                        # 安全打印，避免格式化错误
                        print("[wl] " + str(module['name']) + " " + str(func['name']) + " " + str(phase_name) + " " + str(raw_with_mgmt) + " " + str(wl))

                trace['phases'][phase_name] = {
                    'flow_coeff': flow,
                    'uses_tech_stack': ph['uses_tech_stack'],
                    'raw': round(raw_with_mgmt, 4),
                    'workload': wl
                }

                # 累加原始值（未取整），用于阶段合计的精确计算
                phase_raw_sums[phase_name] += raw_with_mgmt

                if wl > 0:
                    items.append({
                        'phase':    phase_name,
                        'module':   module['name'],
                        'function': func['name'],
                        'workload': wl,
                        'complexity': cx,
                    })

            traces.append(trace)

    # 阶段合计：对原始累加值取整到0.5（比逐项取整更精确）
    phase_totals = OrderedDict()
    for ph in PHASES_V2:
        phase_totals[ph['name']] = round_phase_total(phase_raw_sums[ph['name']])

    # 投产上线 = (需求+UI设计+技术设计+开发+技术测试+性能测试) × go_pct
    go_live_base = sum(phase_totals.get(n, 0) for n in ['需求','UI设计','技术设计','开发','技术测试','性能测试'])
    go_live = round(go_live_base * go_pct, 2)
    phase_totals['投产上线'] = go_live

    total_days = round(sum(phase_totals.values()), 2)

    # 团队成本
    rates = params.get('daily_rates', DEFAULT_PARAMS['daily_rates'])
    team_costs = calculate_team_costs(phase_totals, rates, mgmt, go_live_base)

    # 合规校验
    compliance = validate_compliance(phase_totals, total_days)

    return {
        'items':        items,
        'phase_totals': {k: round(v, 2) for k, v in phase_totals.items()},
        'total_days':   total_days,
        'total_months': round(total_days / 21.75, 2),
        'team_costs':   team_costs,
        'compliance':   compliance,
        'traces':       traces,
    }


def calculate_team_costs(phase_totals, rates, mgmt_coeff, go_live_base):
    """
    产品团队  = 需求人天 × 产品经理单价
    UI团队    = UI设计人天 × UI设计单价
    研发团队  = (技术设计+开发) × 0.4 × 前端单价 + (技术设计+开发) × 0.6 × 后端单价
    测试团队  = 技术测试人天 × 功能测试单价 + 性能测试人天 × 性能测试单价
    项目管理  = (需求+设计+研发+技测+性测) × 管理系数 + 投产上线人天, × PM单价
    """
    req    = phase_totals.get('需求', 0)
    uid    = phase_totals.get('UI设计', 0)
    techd  = phase_totals.get('技术设计', 0)
    dev    = phase_totals.get('开发', 0)
    techu  = phase_totals.get('技术测试', 0)
    perf   = phase_totals.get('性能测试', 0)
    golive = phase_totals.get('投产上线', 0)

    product_cost = round(req * rates.get('product_manager', 2000), 2)
    ui_cost      = round(uid * rates.get('ui_designer', 1800), 2)

    design_dev_total = techd + dev
    frontend_days = round(design_dev_total * 0.4, 2)
    backend_days  = round(design_dev_total * 0.6, 2)
    dev_cost      = round(frontend_days * rates.get('frontend_dev', 1800) + backend_days * rates.get('backend_dev', 2000), 2)

    test_cost = round(techu * rates.get('func_tester', 1500) + perf * rates.get('perf_tester', 2000), 2)

    pm_days = round(go_live_base * mgmt_coeff + golive, 2)
    pm_cost = round(pm_days * rates.get('project_manager', 2000), 2)

    total_cost = round(product_cost + ui_cost + dev_cost + test_cost + pm_cost, 2)

    return {
        'product':  {'days': round(req, 2),              'cost': product_cost,  'label': '产品团队'},
        'ui':       {'days': round(uid, 2),              'cost': ui_cost,       'label': 'UI团队'},
        'dev':      {'days': round(design_dev_total, 2), 'cost': dev_cost,      'label': '研发团队',
                     'frontend_days': frontend_days, 'backend_days': backend_days},
        'testing':  {'days': round(techu + perf, 2),     'cost': test_cost,     'label': '测试团队'},
        'pm':       {'days': pm_days,                    'cost': pm_cost,       'label': '项目管理',
                     'mgmt_coeff': mgmt_coeff},
        'total_cost': total_cost,
    }


def validate_compliance(phase_totals, total_days):
    """
    校验各阶段占比是否在合规区间内。
    注意：合规区间仅覆盖6类阶段（需求/设计/开发/技测/性测/投产），
    分母应使用这6类阶段的合计，而非包含业务测试和准生产的全部总人天。
    """
    design_days = phase_totals.get('UI设计', 0) + phase_totals.get('技术设计', 0)
    check_map = OrderedDict([
        ('需求',    phase_totals.get('需求', 0)),
        ('设计',    design_days),
        ('开发',    phase_totals.get('开发', 0)),
        ('技术测试', phase_totals.get('技术测试', 0)),
        ('性能测试', phase_totals.get('性能测试', 0)),
        ('投产上线', phase_totals.get('投产上线', 0)),
    ])

    # 合规校验分母 = 仅被校验阶段的合计（不含业务测试、准生产）
    check_total = sum(check_map.values())

    results = {}
    all_pass = True
    for phase_label, days in check_map.items():
        lo, hi = COMPLIANCE_RANGES.get(phase_label, (0, 100))
        pct = round(days / check_total * 100, 1) if check_total > 0 else 0
        passed = lo <= pct <= hi
        if not passed:
            all_pass = False
        results[phase_label] = {
            'days': round(days, 2),
            'pct':  pct,
            'min':  lo,
            'max':  hi,
            'pass': passed,
        }

    return {'all_pass': all_pass, 'details': results}

# ===================== 成本消耗预估 =====================

def estimate_cost_consumption(project_data):
    """
    成本消耗预估模块
    根据项目计划和历史数据，预估项目的成本消耗情况
    """
    # 模拟成本消耗预估逻辑
    estimated_cost = project_data.get('estimated_cost', 0)
    start_date = project_data.get('start_date')
    end_date = project_data.get('end_date')
    progress = project_data.get('progress', 0)

    # 计算项目总天数
    import datetime
    start = datetime.datetime.strptime(start_date, '%Y-%m-%d')
    end = datetime.datetime.strptime(end_date, '%Y-%m-%d')
    total_days = (end - start).days
    
    # 计算已过去的天数
    today = datetime.datetime.now()
    elapsed_days = (today - start).days if today > start else 0
    
    # 计算预计成本消耗
    estimated_consumption = estimated_cost * (progress / 100)
    
    # 计算成本消耗趋势
    consumption_trend = []
    for i in range(0, 101, 10):
        consumption_trend.append({
            'progress': i,
            'cost': estimated_cost * (i / 100)
        })
    
    return {
        'estimated_total_cost': estimated_cost,
        'estimated_consumption': estimated_consumption,
        'total_days': total_days,
        'elapsed_days': elapsed_days,
        'consumption_trend': consumption_trend,
        'progress': progress
    }

# ===================== 成本偏差分析 =====================

def analyze_cost_deviation(project_data):
    """
    成本偏差分析模块
    分析项目实际成本与预估成本的偏差情况
    """
    estimated_cost = project_data.get('estimated_cost', 0)
    actual_cost = project_data.get('actual_cost', 0)
    progress = project_data.get('progress', 0)
    
    # 计算成本偏差
    cost_deviation = actual_cost - estimated_cost * (progress / 100)
    cost_deviation_percent = (cost_deviation / (estimated_cost * (progress / 100))) * 100 if progress > 0 else 0
    
    # 分析偏差原因（模拟）
    deviation_reasons = []
    if cost_deviation > 0:
        if cost_deviation_percent > 20:
            deviation_reasons.append('需求变更频繁，导致工作量增加')
            deviation_reasons.append('技术难度超出预期，需要更多资源')
        elif cost_deviation_percent > 10:
            deviation_reasons.append('项目进度延迟，导致人工成本增加')
            deviation_reasons.append('采购成本高于预算')
        else:
            deviation_reasons.append('正常范围的成本波动')
    elif cost_deviation < 0:
        deviation_reasons.append('项目执行效率高于预期')
        deviation_reasons.append('部分需求被简化或取消')
    else:
        deviation_reasons.append('成本控制良好，与预估一致')
    
    # 计算趋势预测
    predicted_final_cost = actual_cost / (progress / 100) if progress > 0 else estimated_cost
    predicted_deviation = predicted_final_cost - estimated_cost
    predicted_deviation_percent = (predicted_deviation / estimated_cost) * 100
    
    return {
        'estimated_cost': estimated_cost,
        'actual_cost': actual_cost,
        'progress': progress,
        'cost_deviation': cost_deviation,
        'cost_deviation_percent': cost_deviation_percent,
        'deviation_reasons': deviation_reasons,
        'predicted_final_cost': predicted_final_cost,
        'predicted_deviation': predicted_deviation,
        'predicted_deviation_percent': predicted_deviation_percent
    }

# ===================== Excel 生成 =====================

def make_border(style='thin'):
    s = Side(style=style)
    return Border(left=s, right=s, top=s, bottom=s)


def set_range(ws, r1, c1, r2, c2, value, bold=False, sz=10, ha='center', va='center',
              bg=None, fc='000000', wrap=False, num_fmt=None):
    if r1 != r2 or c1 != c2:
        ws.merge_cells(start_row=r1, start_column=c1, end_row=r2, end_column=c2)
    cell = ws.cell(row=r1, column=c1)
    cell.value = value
    cell.font = Font(name='Arial', size=sz, bold=bold, color=fc)
    cell.alignment = Alignment(horizontal=ha, vertical=va, wrap_text=wrap)
    if num_fmt:
        cell.number_format = num_fmt
    bd = make_border()
    for row in range(r1, r2 + 1):
        for col in range(c1, c2 + 1):
            c = ws.cell(row=row, column=col)
            c.border = bd
            if bg:
                c.fill = PatternFill('solid', start_color=bg)
    return cell


def generate_excel(calc_result, project_name='', project_id='', system_name='', params=None):
    if params is None:
        params = DEFAULT_PARAMS

    items         = calc_result['items']
    phase_totals  = calc_result['phase_totals']
    total_days    = calc_result['total_days']
    total_months  = calc_result['total_months']
    team_costs    = calc_result['team_costs']
    compliance    = calc_result['compliance']
    traces        = calc_result.get('traces', [])

    wb  = Workbook()
    HBG = 'BDD7EE'
    IBG = 'D9E1F2'

    # ── Sheet 1: 填写说明 ──
    ws1 = wb.active
    ws1.title = '填写说明'
    ws1.column_dimensions['B'].width = 90
    set_range(ws1, 2, 2, 2, 13, '表格说明', bold=True, sz=12, bg='4472C4', fc='FFFFFF')
    ws1.row_dimensions[2].height = 28
    note = (
        "1. 项目工作量评估表：项目类型涉及实施类项目时需要填写的工作量，工作量内容仅为科技侧工作量，不含业务侧工作量。\n"
        "2. 成本汇总表：各团队（产品/UI/研发/测试/项目管理）的人天与成本汇总。\n"
        "3. 合规校验表：各阶段工作量占比是否符合行业标准区间。\n"
        "4. 产品采购评估表：项目类型涉及产品项目时需要填写。\n"
        "5. 计算参数：本次评估使用的所有系数参数，支持追溯。"
    )
    set_range(ws1, 3, 2, 7, 13, note, ha='left', va='top', wrap=True)
    for r in range(3, 8):
        ws1.row_dimensions[r].height = 20

    # ── Sheet 2: 项目工作量评估表 ──
    ws2 = wb.create_sheet('项目工作量评估表')
    col_w = {1:2, 2:10, 3:20, 4:22, 5:4, 6:30, 7:4, 8:4, 9:4, 10:4, 11:4, 12:8, 13:8, 14:8, 15:30}
    for c, w in col_w.items():
        ws2.column_dimensions[get_column_letter(c)].width = w

    set_range(ws2, 2, 2, 2, 15, '项目工作量评估表（仅为科技侧工作量）',
              bold=True, sz=14, bg='4472C4', fc='FFFFFF')
    ws2.row_dimensions[2].height = 34

    set_range(ws2, 3, 2, 3, 4, f'项目编号：{project_id}',     ha='left', bg=IBG)
    set_range(ws2, 3, 5, 3, 11, f'项目名称：{project_name}',  ha='left', bg=IBG)
    set_range(ws2, 3, 12, 3, 15, f'系统名称：{system_name}',  ha='left', bg=IBG)
    ws2.row_dimensions[3].height = 22

    for c1, c2, txt in [
        (2,2,'项目阶段'),(3,3,'系统模块'),(4,5,'功能'),
        (6,11,'工作项描述'),(12,12,'备注'),
        (13,13,'工作量\n(人天)'),(14,14,'合计\n(人天)'),(15,15,'任务描述')
    ]:
        set_range(ws2, 4, c1, 4, c2, txt, bold=True, bg=HBG, wrap=True)
    ws2.row_dimensions[4].height = 34

    # 组织数据
    organized = OrderedDict()
    for item in items:
        p, m = item['phase'], item['module']
        if p not in organized:
            organized[p] = OrderedDict()
        if m not in organized[p]:
            organized[p][m] = []
        organized[p][m].append(item)

    PHASE_DESC = {
        '需求':    '根据需求文档，配合业务开展技术调研、需求评审。',
        'UI设计':  '根据业务需求进行交互设计、视觉稿输出及评审。',
        '技术设计': '根据业务需求进行详细技术设计、数据模型设计。',
        '开发':    '根据详细设计进行编码开发，含代码审查及单元测试。',
        '技术测试': '编写技术测试方案，执行测试，提交缺陷并跟踪闭环。',
        '性能测试': '制定性能测试方案，执行压力与负载测试，分析优化。',
        '业务测试': '协助业务人员制定测试方案，执行业务功能验证。',
        '准生产':   '准生产环境部署验证，功能回归测试，问题排查。',
        '投产上线': '上线部署，生产环境功能验证，提供上线技术支持。',
    }

    row = 5
    for phase_name, modules_dict in organized.items():
        bg = PHASE_BG.get(phase_name, 'FFFFFF')
        all_items_in_phase = [i for ml in modules_dict.values() for i in ml]
        ph_start = row

        first_in_module = {}
        for item in all_items_in_phase:
            if item['module'] not in first_in_module:
                first_in_module[item['module']] = item

        for item in all_items_in_phase:
            ws2.row_dimensions[row].height = 18
            c = ws2.cell(row=row, column=3)
            if item is first_in_module[item['module']]:
                c.value = item['module']
            c.font      = Font(name='Arial', size=9)
            c.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
            c.border    = make_border()
            c.fill      = PatternFill('solid', start_color=bg)

            set_range(ws2, row, 4, row, 5,  item['function'],  ha='left', bg=bg, sz=9, wrap=True)
            fixed_desc = PHASE_DESC.get(phase_name, '')
            work_item_desc = f"{item['function']}：{fixed_desc}" if fixed_desc else item['function']
            set_range(ws2, row, 6, row, 11, work_item_desc, ha='left', bg=bg, sz=9, wrap=True)
            set_range(ws2, row, 12, row, 12, '', bg=bg)

            c = ws2.cell(row=row, column=13)
            c.value      = item['workload']
            c.font       = Font(name='Arial', size=9)
            c.alignment  = Alignment(horizontal='center', vertical='center')
            c.border     = make_border()
            c.fill       = PatternFill('solid', start_color=bg)
            c.number_format = '#,##0.00'
            row += 1

        ph_end = row - 1
        set_range(ws2, ph_start, 2, ph_end, 2, phase_name, bold=True, bg=bg)
        stc = set_range(ws2, ph_start, 14, ph_end, 14,
                        f'=SUM(M{ph_start}:M{ph_end})', bold=True, bg=bg)
        stc.number_format = '#,##0.00'
        set_range(ws2, ph_start, 15, ph_end, 15, PHASE_DESC.get(phase_name, ''), ha='left', va='top', bg=bg, sz=9, wrap=True)

    # 追加投产上线（固定行，来自calc_result）
    golive_bg = PHASE_BG.get('投产上线', 'F3E5F5')
    set_range(ws2, row, 2, row, 13, f'投产上线（以上阶段合计 × {int(params.get("go_live_pct",0.02)*100)}%）',
              bold=True, bg=golive_bg, ha='left')
    golive_cell = ws2.cell(row=row, column=14)
    golive_cell.value  = phase_totals.get('投产上线', 0)
    golive_cell.font   = Font(name='Arial', size=10, bold=True)
    golive_cell.alignment = Alignment(horizontal='center', vertical='center')
    golive_cell.border = make_border()
    golive_cell.fill   = PatternFill('solid', start_color=golive_bg)
    golive_cell.number_format = '#,##0.00'
    set_range(ws2, row, 15, row, 15, '上线部署及技术支持', bg=golive_bg)
    row += 1

    total_row = row
    set_range(ws2, total_row, 2, total_row, 3, '总计（人天）', bold=True, sz=11, bg='4472C4', fc='FFFFFF')
    tc = set_range(ws2, total_row, 4, total_row, 15, total_days, bold=True, sz=11, bg='4472C4', fc='FFFFFF')
    tc.number_format = '#,##0.00'
    ws2.row_dimensions[total_row].height = 24

    month_row = total_row + 1
    set_range(ws2, month_row, 2, month_row, 3, '总计（人月）', bold=True, sz=11, bg='4472C4', fc='FFFFFF')
    mc = set_range(ws2, month_row, 4, month_row, 15, total_months, bold=True, sz=11, bg='4472C4', fc='FFFFFF')
    mc.number_format = '#,##0.00'
    ws2.row_dimensions[month_row].height = 24

    # ── Sheet 3: 成本汇总表 ──
    ws3 = wb.create_sheet('成本汇总表')
    ws3.column_dimensions['B'].width = 18
    ws3.column_dimensions['C'].width = 14
    ws3.column_dimensions['D'].width = 14
    ws3.column_dimensions['E'].width = 20
    set_range(ws3, 1, 2, 1, 6, '项目团队成本汇总表', bold=True, sz=13, bg='4472C4', fc='FFFFFF')
    ws3.row_dimensions[1].height = 30
    set_range(ws3, 2, 2, 2, 6, f'项目：{project_name}  |  系统：{system_name}', ha='left', sz=10, bg=IBG)
    for c1, c2, txt in [(2,2,'团队'), (3,3,'工作量（人天）'), (4,4,'单价（元/天）'), (5,5,'成本（元）'), (6,6,'说明')]:
        set_range(ws3, 3, c1, 3, c2, txt, bold=True, bg=HBG)
    ws3.row_dimensions[3].height = 22

    cost_rows = [
        ('产品团队',  team_costs['product']['days'], params['daily_rates']['product_manager'],   team_costs['product']['cost'],  '需求分析阶段'),
        ('UI团队',    team_costs['ui']['days'],       params['daily_rates']['ui_designer'],       team_costs['ui']['cost'],       'UI/UX设计阶段'),
        ('研发团队',  team_costs['dev']['days'],      '前端/后端混合',                           team_costs['dev']['cost'],      f"前端{team_costs['dev']['frontend_days']}天+后端{team_costs['dev']['backend_days']}天"),
        ('测试团队',  team_costs['testing']['days'],  '功能/性能混合',                           team_costs['testing']['cost'],  '技术测试+性能测试'),
        ('项目管理',  team_costs['pm']['days'],       params['daily_rates']['project_manager'],   team_costs['pm']['cost'],       f"管理系数{int(team_costs['pm']['mgmt_coeff']*100)}%+投产"),
    ]
    for i, (label, days, rate, cost, note) in enumerate(cost_rows, start=4):
        bg = 'F8FAFC' if i % 2 == 0 else 'FFFFFF'
        set_range(ws3, i, 2, i, 2, label,  bold=True,  bg=bg, ha='left')
        set_range(ws3, i, 3, i, 3, round(days, 2), bg=bg, num_fmt='#,##0.00')
        set_range(ws3, i, 4, i, 4, rate,   bg=bg, ha='left')
        set_range(ws3, i, 5, i, 5, round(cost, 2), bg=bg, num_fmt='#,##0.00')
        set_range(ws3, i, 6, i, 6, note,   bg=bg, ha='left', sz=9)
        ws3.row_dimensions[i].height = 20

    total_cost_row = len(cost_rows) + 4
    set_range(ws3, total_cost_row, 2, total_cost_row, 4, '项目总成本', bold=True, sz=11, bg='4472C4', fc='FFFFFF')
    tc3 = set_range(ws3, total_cost_row, 5, total_cost_row, 6, team_costs['total_cost'],
                    bold=True, sz=11, bg='4472C4', fc='FFFFFF')
    tc3.number_format = '#,##0.00'
    ws3.row_dimensions[total_cost_row].height = 26

    # ── Sheet 4: 合规校验 ──
    ws4 = wb.create_sheet('合规校验报告')
    ws4.column_dimensions['B'].width = 14
    ws4.column_dimensions['C'].width = 12
    ws4.column_dimensions['D'].width = 10
    ws4.column_dimensions['E'].width = 10
    ws4.column_dimensions['F'].width = 10
    ws4.column_dimensions['G'].width = 12
    ws4.column_dimensions['H'].width = 16
    set_range(ws4, 1, 2, 1, 8, '各阶段工作量合规校验报告', bold=True, sz=13, bg='4472C4', fc='FFFFFF')
    ws4.row_dimensions[1].height = 30
    overall = '✅ 全部通过' if compliance['all_pass'] else '⚠️ 存在超出区间'
    ovbg = '4CAF50' if compliance['all_pass'] else 'FF9800'
    set_range(ws4, 2, 2, 2, 8, f'总体结果：{overall}  （总计 {total_days} 人天 / {total_months} 人月）',
              ha='left', sz=10, bg=ovbg, fc='FFFFFF', bold=True)
    for c1, c2, txt in [(2,2,'阶段'), (3,3,'人天'), (4,4,'占比%'), (5,5,'最低%'), (6,6,'最高%'), (7,7,'结果'), (8,8,'备注')]:
        set_range(ws4, 3, c1, 3, c2, txt, bold=True, bg=HBG)
    ws4.row_dimensions[3].height = 22
    for i, (phase, detail) in enumerate(compliance['details'].items(), start=4):
        passed = detail['pass']
        bg = 'F1F8E9' if passed else 'FFF3E0'
        result_txt = '✅ 通过' if passed else '⚠️ 超出区间'
        note_txt = '' if passed else f"实际{detail['pct']}%，区间[{detail['min']}%,{detail['max']}%]"
        set_range(ws4, i, 2, i, 2, phase,             bold=True, bg=bg, ha='left')
        set_range(ws4, i, 3, i, 3, detail['days'],    bg=bg, num_fmt='#,##0.00')
        set_range(ws4, i, 4, i, 4, detail['pct'],     bg=bg, num_fmt='0.0')
        set_range(ws4, i, 5, i, 5, detail['min'],     bg=bg)
        set_range(ws4, i, 6, i, 6, detail['max'],     bg=bg)
        set_range(ws4, i, 7, i, 7, result_txt,        bg=bg, fc='166534' if passed else 'B45309')
        set_range(ws4, i, 8, i, 8, note_txt,          bg=bg, ha='left', sz=9)
        ws4.row_dimensions[i].height = 20

    # ── Sheet 5: 计算参数 ──
    ws5 = wb.create_sheet('计算参数')
    ws5.column_dimensions['B'].width = 22
    ws5.column_dimensions['C'].width = 18
    ws5.column_dimensions['D'].width = 28
    set_range(ws5, 1, 2, 1, 4, '本次评估计算参数', bold=True, sz=13, bg='4472C4', fc='FFFFFF')
    ws5.row_dimensions[1].height = 28
    param_rows = [
        ('系统关联度系数', params.get('association_coeff', 1.0),       '关联系统数量决定取值(1/1.5/2/3)'),
        ('技术栈难度系数', params.get('tech_stack_coeff', 1.2),        '常规1.0/微服务1.2/AI中台1.4/分布式1.6'),
        ('管理系数',       params.get('mgmt_coeff', 0.15),             '小型团队≤10人: 0.15; 中型0.20; 大型0.25'),
        ('投产上线比例',   f"{int(params.get('go_live_pct', 0.02)*100)}%", '前置阶段合计 × 该系数'),
    ]
    for ph_key, ph_info in [
        ('requirements', '需求流程系数'), ('ui_design', 'UI设计流程系数'),
        ('tech_design', '技术设计流程系数'), ('development', '开发流程系数'),
        ('tech_testing', '技术测试流程系数'), ('perf_testing', '性能测试流程系数'),
    ]:
        param_rows.append((ph_info, params['flow_coeffs'].get(ph_key, '-'), ''))
    for cb_key, cb_label in [
        ('very_basic', '较为基础复杂度基准'),
        ('basic', '基础复杂度基准'),
        ('medium', '中等复杂度基准'),
        ('complex', '复杂复杂度基准'),
        ('very_complex', '极复杂复杂度基准')
    ]:
        param_rows.append((cb_label, params['complexity_base'].get(cb_key, '-'), '人天'))
    for role, label in [
        ('product_manager','产品经理单价'), ('ui_designer','UI设计单价'),
        ('frontend_dev','前端开发单价'), ('backend_dev','后端开发单价'),
        ('func_tester','功能测试单价'), ('perf_tester','性能测试单价'), ('project_manager','项目经理单价')
    ]:
        param_rows.append((label, params['daily_rates'].get(role, '-'), '元/人天'))

    for c1, c2, txt in [(2,2,'参数名称'), (3,3,'参数值'), (4,4,'说明')]:
        set_range(ws5, 2, c1, 2, c2, txt, bold=True, bg=HBG)
    for i, (name, val, note) in enumerate(param_rows, start=3):
        bg = 'F8FAFC' if i % 2 == 0 else 'FFFFFF'
        set_range(ws5, i, 2, i, 2, name, ha='left', bg=bg)
        set_range(ws5, i, 3, i, 3, val,  bg=bg)
        set_range(ws5, i, 4, i, 4, note, ha='left', bg=bg, sz=9)
        ws5.row_dimensions[i].height = 18

    out = io.BytesIO()
    wb.save(out)
    out.seek(0)
    return out.getvalue()

# ===================== Tornado 处理器 =====================

class MainHandler(tornado.web.RequestHandler):
    def get(self):
        html_path = Path(__file__).parent / 'index.html'
        with open(html_path, 'r', encoding='utf-8') as f:
            self.write(f.read())


class LoginHandler(tornado.web.RequestHandler):
    def post(self):
        self.set_header('Access-Control-Allow-Origin', '*')
        self.set_header('Content-Type', 'application/json; charset=utf-8')
        try:
            data = json.loads(self.request.body.decode('utf-8'))
            username = data.get('username')
            password = data.get('password')
            
            if username in USERS and USERS[username]['password'] == password:
                self.write(json.dumps({
                    'success': True,
                    'user': {
                        'username': username,
                        'role': USERS[username]['role']
                    }
                }, ensure_ascii=False))
            else:
                self.set_status(401)
                self.write(json.dumps({
                    'success': False,
                    'error': '用户名或密码错误'
                }, ensure_ascii=False))
        except Exception as e:
            self.set_status(500)
            self.write(json.dumps({
                'success': False,
                'error': str(e)
            }, ensure_ascii=False))


class ProjectsHandler(tornado.web.RequestHandler):
    def get(self):
        self.set_header('Access-Control-Allow-Origin', '*')
        self.set_header('Content-Type', 'application/json; charset=utf-8')
        self.write(json.dumps(PROJECTS, ensure_ascii=False))


class ProjectDetailHandler(tornado.web.RequestHandler):
    def get(self, project_id):
        self.set_header('Access-Control-Allow-Origin', '*')
        self.set_header('Content-Type', 'application/json; charset=utf-8')
        project = next((p for p in PROJECTS if p['id'] == project_id), None)
        if project:
            self.write(json.dumps(project, ensure_ascii=False))
        else:
            self.set_status(404)
            self.write(json.dumps({'error': '项目不存在'}, ensure_ascii=False))


class CostEstimationHandler(tornado.web.RequestHandler):
    def post(self, project_id):
        self.set_header('Access-Control-Allow-Origin', '*')
        self.set_header('Content-Type', 'application/json; charset=utf-8')
        try:
            project = next((p for p in PROJECTS if p['id'] == project_id), None)
            if not project:
                self.set_status(404)
                self.write(json.dumps({'error': '项目不存在'}, ensure_ascii=False))
                return
            
            result = estimate_cost_consumption(project)
            self.write(json.dumps(result, ensure_ascii=False))
        except Exception as e:
            self.set_status(500)
            self.write(json.dumps({'error': str(e)}, ensure_ascii=False))


class CostDeviationHandler(tornado.web.RequestHandler):
    def post(self, project_id):
        self.set_header('Access-Control-Allow-Origin', '*')
        self.set_header('Content-Type', 'application/json; charset=utf-8')
        try:
            project = next((p for p in PROJECTS if p['id'] == project_id), None)
            if not project:
                self.set_status(404)
                self.write(json.dumps({'error': '项目不存在'}, ensure_ascii=False))
                return
            
            result = analyze_cost_deviation(project)
            self.write(json.dumps(result, ensure_ascii=False))
        except Exception as e:
            self.set_status(500)
            self.write(json.dumps({'error': str(e)}, ensure_ascii=False))


class AnalyzeHandler(tornado.web.RequestHandler):
    def post(self):
        self.set_header('Access-Control-Allow-Origin', '*')
        self.set_header('Content-Type', 'application/json; charset=utf-8')
        try:
            doc_text = self.get_argument('doc_text', '')
            if 'file' in self.request.files:
                finfo    = self.request.files['file'][0]
                doc_text = extract_text(finfo['body'], finfo['filename'])
            if not doc_text.strip():
                self.set_status(400)
                self.write(json.dumps({'error': '无法读取文档内容，请确认文件格式或直接粘贴文本'}, ensure_ascii=False))
                return
            result = analyze_requirements(doc_text, API_KEY)
            if not result.get('modules'):
                self.write(json.dumps({
                    'error': 'AI分析完成但未识别到功能模块，请尝试直接粘贴需求文本。',
                    **result
                }, ensure_ascii=False))
                return
            # 统计功能点总数，超过30个给出警告
            total_funcs = sum(len(m.get('functions', [])) for m in result.get('modules', []))
            if total_funcs > 30:
                result['warning'] = (
                    f'当前识别出 {total_funcs} 个功能点，超过建议上限（30个）。'
                    f'由于模型能力有限，功能点过多可能导致评估精度下降，建议精简或分批评估。'
                )
            self.write(json.dumps(result, ensure_ascii=False))
        except Exception as e:
            self.set_status(500)
            print(traceback.format_exc())
            self.write(json.dumps({'error': str(e)}, ensure_ascii=False))


class EstimateHandler(tornado.web.RequestHandler):
    def post(self):
        self.set_header('Access-Control-Allow-Origin', '*')
        self.set_header('Content-Type', 'application/json; charset=utf-8')
        try:
            data    = json.loads(self.request.body.decode('utf-8'))
            modules = data.get('modules', [])
            params  = _merge_params(data.get('params'))
            result  = estimate_workload_v2(modules, params)
            self.write(json.dumps({
                'total_items':  len(result['items']),
                'total_days':   result['total_days'],
                'total_months': result['total_months'],
                'phase_totals': result['phase_totals'],
                'team_costs':   result['team_costs'],
                'compliance':   result['compliance'],
                'traces':       result.get('traces', []),
            }, ensure_ascii=False))
        except Exception as e:
            self.set_status(500)
            self.write(json.dumps({'error': str(e)}, ensure_ascii=False))


class GenerateHandler(tornado.web.RequestHandler):
    def post(self):
        self.set_header('Access-Control-Allow-Origin', '*')
        try:
            data         = json.loads(self.request.body.decode('utf-8'))
            modules      = data.get('modules', [])
            project_name = data.get('project_name', '')
            project_id   = data.get('project_id', '')
            system_name  = data.get('system_name', '')
            params       = _merge_params(data.get('params'))

            if not modules:
                self.set_status(400)
                self.set_header('Content-Type', 'application/json; charset=utf-8')
                self.write(json.dumps({'error': '缺少功能模块数据'}, ensure_ascii=False))
                return

            calc_result  = estimate_workload_v2(modules, params)
            excel_bytes  = generate_excel(calc_result, project_name, project_id, system_name, params)

            ts    = time.strftime('%Y%m%d_%H%M%S')
            fname = f'工作量评估表_{project_name}_{ts}.xlsx' if project_name else f'工作量评估表_{ts}.xlsx'
            (OUTPUT_DIR / fname).write_bytes(excel_bytes)

            from urllib.parse import quote
            enc = quote(fname)
            self.set_header('Content-Type', 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
            self.set_header('Content-Disposition', f"attachment; filename*=UTF-8''{enc}")
            self.set_header('X-Total-Days',    str(calc_result['total_days']))
            self.set_header('X-Total-Months',  str(calc_result['total_months']))
            self.set_header('X-Total-Cost',    str(calc_result['team_costs']['total_cost']))
            self.set_header('X-Total-Items',   str(len(calc_result['items'])))
            self.set_header('X-Filename',      enc)
            self.set_header('X-All-Pass',      str(calc_result['compliance']['all_pass']))
            self.write(excel_bytes)
        except Exception as e:
            self.set_status(500)
            self.set_header('Content-Type', 'application/json; charset=utf-8')
            self.write(json.dumps({'error': str(e), 'trace': traceback.format_exc()}, ensure_ascii=False))


def _merge_params(user_params):
    """将用户参数深度合并到DEFAULT_PARAMS中"""
    import copy
    p = copy.deepcopy(DEFAULT_PARAMS)
    if not user_params:
        return p
    for key in ('association_coeff', 'tech_stack_coeff', 'mgmt_coeff', 'go_live_pct'):
        if key in user_params:
            p[key] = float(user_params[key])
    if 'flow_coeffs' in user_params:
        p['flow_coeffs'].update({k: float(v) for k, v in user_params['flow_coeffs'].items()})
    if 'complexity_base' in user_params:
        p['complexity_base'].update({k: float(v) for k, v in user_params['complexity_base'].items()})
    if 'daily_rates' in user_params:
        p['daily_rates'].update({k: float(v) for k, v in user_params['daily_rates'].items()})
    return p


def make_app():
    return tornado.web.Application([
        (r'/',             MainHandler),
        (r'/api/login',    LoginHandler),
        (r'/api/projects', ProjectsHandler),
        (r'/api/project/(.*)', ProjectDetailHandler),
        (r'/api/project/(.*)/cost-estimation', CostEstimationHandler),
        (r'/api/project/(.*)/cost-deviation', CostDeviationHandler),
        (r'/api/analyze',  AnalyzeHandler),
        (r'/api/estimate', EstimateHandler),
        (r'/api/generate', GenerateHandler),
    ])


if __name__ == '__main__':
    print("Starting server...")
    try:
        import asyncio
        print("Setting up event loop...")
        asyncio.set_event_loop(asyncio.new_event_loop())
        print("Creating application...")
        app = make_app()
        print(f"Listening on port {PORT}...")
        app.listen(PORT)
        print(f'\n项目成本估算监控系统 v1.0 已启动')
        print(f'访问地址: http://localhost:{PORT}')
        print(f'输出目录: {OUTPUT_DIR}')
        print(f'\n按 Ctrl+C 停止服务\n')
        print("Starting IOLoop...")
        tornado.ioloop.IOLoop.current().start()
    except Exception as e:
        print(f"Error starting server: {e}")
        import traceback
        traceback.print_exc()